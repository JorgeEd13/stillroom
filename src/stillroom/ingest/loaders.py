# Copyright (c) Jorge Ribeiro.
# Licensed under PolyForm Shield 1.0.0 - see LICENSE.md.
# Source-available, not open source: any purpose is permitted except
# providing a product that competes with this software.
"""Read a client's documents into plain text.

`receivables-agent` reads exactly one Markdown file that was *written to be
chunkable*. A client's corpus is the opposite: whatever happened to accumulate
in a folder, in four formats, written by people who were not thinking about
retrieval. So this module's job is narrower and duller than it looks — turn
bytes into text plus a stable identity, and refuse clearly when it cannot.

**No OCR, deliberately**. A PDF that is a photograph of a page yields
no text, and the honest response is to name the file and skip it rather than
index an empty document that will never be retrieved and never be missed. The
intake checklist asks up front whether any documents
are scanned images, so this path should be rare — but "rare" is not "never",
and a silent empty document is a bug the client discovers as a wrong answer
weeks later.
"""

from __future__ import annotations

import codecs
import logging
import unicodedata
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class RawDocument:
    """One source document, as text.

    `source` is the path *relative to the corpus root*, not the absolute path.
    Citations show it to the client, so it must be something they recognise —
    and it must not leak the directory layout of the machine it was built on.
    """

    source: str
    text: str


class UnreadableDocument(Exception):
    """Raised when a file matched the include list but yielded no usable text."""


# Byte-order marks, longest first: a UTF-32 LE mark begins with the same two
# bytes as a UTF-16 LE one, so testing in the other order silently mis-decodes.
_BOMS: tuple[tuple[bytes, str], ...] = (
    (codecs.BOM_UTF32_LE, "utf-32"),
    (codecs.BOM_UTF32_BE, "utf-32"),
    (codecs.BOM_UTF8, "utf-8-sig"),
    (codecs.BOM_UTF16_LE, "utf-16"),
    (codecs.BOM_UTF16_BE, "utf-16"),
)

# ⚠️ The candidate list is a **business** decision, not a technical one, and
# leaving it off got the answer wrong. Unrestricted, the detector read a short
# Brazilian invoice as **cp1250** (Central European) and produced `manutençăo`
# for `manutenção` — plausible-looking, subtly wrong, and unsearchable. Given
# the same bytes and this list it returns cp1252 and the right word.
#
# The list says which languages this product is built for (English and Brazilian
# Portuguese), so a corpus from outside them needs `[corpus] encoding` set
# explicitly — which is the honest way round: an assumption in a config file
# somebody chose, rather than one buried in a detector's ranking.
_WESTERN_CODEPAGES = ["utf_8", "cp1252", "latin_1", "mac_roman"]

# Control characters that separate words in the systems that emit them —
# vertical tab, form feed, next line. Deleting these outright glues the words
# either side together: `Control\x0bchars` became `Controlchars`, which matches
# no search anybody would type.
_C0_WHITESPACE = "\x0b\x0c\x1c\x1d\x1e\x1f\x85"


def _decode(raw: bytes, override: str | None = None) -> str:
    """Bytes to text, for documents nobody wrote with us in mind.

    ⚠️ **This used to be `read_text(encoding="utf-8", errors="replace")`, and
    that one line quietly destroyed real client documents.** The comment
    defending it — a stray byte must not abort a 200-file ingest — was right
    about the goal and wrong about the mechanism: `replace` does not survive a
    non-UTF-8 file, it *ruins* it. Two cases from a test corpus:

    - **cp1252/latin-1**, which is what a decade of Windows-authored Brazilian
      documents actually are: `manutenção` was indexed as `manuten��o`.
      Every accented word became unsearchable, and the mangled text is what the
      citation panel shows the client, in their own language.
    - **UTF-16**, which is what Notepad's "Unicode" and a PowerShell redirect
      produce: every character came back with a `\\x00` between its bytes. The
      file counted as a **successfully indexed document** and contained nothing
      retrievable — the worst of the two outcomes, because it is silent.

    So: an explicit override if the engagement set one, then the mark, then
    strict UTF-8, then a detector over a deliberate candidate list, then cp1252
    as the floor. `errors="replace"` survives only in that last step, where the
    alternative is refusing a file we could still mostly read.
    """
    if override:
        # An engagement that knows its own corpus beats any amount of guessing.
        return raw.decode(override, errors="replace")

    for mark, encoding in _BOMS:
        if raw.startswith(mark):
            return raw.decode(encoding, errors="replace")

    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass

    try:
        # Already present in the image through `requests`; declared in
        # `pyproject.toml` so the bill of materials names it on purpose rather
        # than inheriting it by accident.
        from charset_normalizer import from_bytes

        best = from_bytes(raw, cp_isolation=_WESTERN_CODEPAGES).best()
        if best is not None:
            return str(best)
    except ImportError:  # pragma: no cover - only if the extra is stripped
        pass

    # cp1252 rather than latin-1: it is a superset over the range that matters
    # (quotes, dashes, the euro sign) and decodes everything latin-1 would.
    return raw.decode("cp1252", errors="replace")


def _clean(text: str) -> str:
    """Remove what is invisible or deceptive, keep what is merely foreign.

    Three classes, and the third is the reason this is not cosmetic:

    - **Zero-width characters** split a word for the retriever while looking
      identical to the client — `warranty​ period` does not match a search
      for "warranty period".
    - **C0 control characters** arrive in exports from old systems and render as
      boxes or nothing at all in the citation panel.
    - **Bidirectional overrides** make displayed text read differently from the
      text that is stored. On a page whose job is to show the client exactly
      what their document says, a character that reverses the rest of the line
      is not a formatting quirk.

    Arabic, Hebrew, CJK and emoji are untouched. The target is the invisible
    formatting characters, never the writing system.

    ⚠️ **Selected by Unicode category `Cf`, not by an explicit list**, because
    the explicit list was wrong the first time: it caught the zero-width space
    and missed `U+200E LEFT-TO-RIGHT MARK`, which sat in the indexed text
    afterwards. A category covers the whole class, including the members nobody
    thought to enumerate.

    The known cost, accepted rather than hidden: `Cf` also contains the
    zero-width non-joiner, which is **meaningful** in Persian and Urdu. Removing
    it there alters words. This product is built for English and Brazilian
    Portuguese; a corpus in a script that needs it would need this reconsidered,
    not a config flag pretending the trade-off is not there.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    for ch in _C0_WHITESPACE:
        text = text.replace(ch, "\n")
    stripped = (
        ch
        for ch in text
        if unicodedata.category(ch) != "Cf"
        and (ch in "\n\t" or unicodedata.category(ch) != "Cc")
    )
    return "".join(stripped)


def _load_text(path: Path, encoding: str | None = None) -> str:
    return _decode(path.read_bytes(), encoding)


def read_plain_text(path: Path, encoding: str | None = None) -> str:
    """Decode and clean one text file, with the corpus rules.

    Public because the standing-context document is read through it too: notes
    the client wrote in Word and saved as `.txt` deserve the same encoding
    ladder as the rest of their corpus, and a second, simpler reader elsewhere
    would be a second place for the cp1252 bug to live.
    """
    return _clean(_decode(path.read_bytes(), encoding))


def _load_pdf(path: Path, encoding: str | None = None) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - install-time failure
        raise UnreadableDocument(
            f"{path.name}: PDF support needs the 'docs' extra (pypdf)."
        ) from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text.strip():
        raise UnreadableDocument(
            f"{path.name}: no extractable text — likely a scanned image. "
            "OCR is out of scope; this document was skipped."
        )
    return text


def _load_docx(path: Path, encoding: str | None = None) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - install-time failure
        raise UnreadableDocument(
            f"{path.name}: Word support needs the 'docs' extra (python-docx)."
        ) from exc

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables carry a lot of the actual policy in real corpora (rate cards,
    # thresholds, escalation matrices) and are invisible to `paragraphs`.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    if not text.strip():
        raise UnreadableDocument(f"{path.name}: contains no text.")
    return text


def _load_csv(path: Path, encoding: str | None = None) -> str:
    """A comma/semicolon-separated table, rendered row-per-line.

    Spreadsheets exported to CSV are one of the commonest things a real company
    has — price lists, inventories, exports from a line-of-business system. Each
    row becomes ``cell | cell``, the same shape a Word table gets, so the model
    sees a coherent record rather than a wall of commas. `csv.Sniffer` picks the
    delimiter, because a European export is as likely to be semicolon-separated.
    """
    import csv
    import io

    text = _decode(path.read_bytes(), encoding)
    sample = text[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel
    rows = []
    for row in csv.reader(io.StringIO(text), dialect):
        cells = [c.strip() for c in row if c and c.strip()]
        if cells:
            rows.append(" | ".join(cells))
    out = "\n".join(rows)
    if not out.strip():
        raise UnreadableDocument(f"{path.name}: contains no rows.")
    return out


def _load_xlsx(path: Path, encoding: str | None = None) -> str:
    """An Excel workbook, one sheet after another, row-per-line.

    ``data_only=True`` reads the last-computed value of a formula cell, not the
    formula — a client's ``=SUM(...)`` total is a fact to answer from, and the
    formula text is noise. ``read_only=True`` streams a large workbook instead of
    building it in memory. The sheet title prefixes its rows so a workbook whose
    tabs are ``Q3`` and ``Q4`` does not silently merge the two.
    """
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover - install-time failure
        raise UnreadableDocument(
            f"{path.name}: Excel support needs the 'docs' extra (openpyxl)."
        ) from exc

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    try:
        parts = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(sheet.title + "\n" + "\n".join(rows))
    finally:
        workbook.close()

    text = "\n\n".join(parts)
    if not text.strip():
        raise UnreadableDocument(f"{path.name}: contains no cells.")
    return text


_LOADERS = {
    ".md": _load_text,
    ".txt": _load_text,
    ".csv": _load_csv,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".xlsx": _load_xlsx,
}


# Files an operating system leaves in a folder by itself. Reporting these as
# "skipped" would bury the two lines a client actually needs to read.
_CLUTTER = frozenset(
    {".DS_Store", "Thumbs.db", "desktop.ini", ".gitkeep", "Icon\r"}
)


def _is_inside(path: Path, root_real: Path) -> bool:
    """Does this file really live under the corpus root?

    Compared after `resolve()` on both sides, so a symlink is judged by where it
    lands rather than where it sits.
    """
    try:
        return path.resolve().is_relative_to(root_real)
    except OSError:
        # A broken or looping link resolves to nothing usable. Not inside.
        return False


def load_document(
    path: Path, root: Path, encoding: str | None = None
) -> RawDocument:
    """Load one file, or raise `UnreadableDocument`.

    `encoding` applies to plain text only; PDF and Word carry their own.
    """
    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        raise UnreadableDocument(f"{path.name}: unsupported file type.")

    try:
        text = loader(path, encoding)
    except UnreadableDocument:
        raise
    except Exception as exc:
        # Deliberately broad. These are third-party parsers pointed at whatever
        # a client happened to put in a folder — a truncated PDF, a file with
        # the wrong extension, a Word document written by software nobody has
        # heard of. They raise from a wide and undocumented set, and every one
        # of them means the same thing here: this file yielded no text. The
        # alternative is an ingest that dies on file 180 of 200.
        # ⚠️ The parser's message often contains the **absolute path**, and this
        # string is printed to the client at handover. `RawDocument.source`
        # already refuses to leak the build machine's directory layout; a skip
        # reason must not undo that. The path is replaced rather than
        # dropped because the rest of the parser's complaint is the only clue
        # about what is wrong with the file.
        detail = str(exc).replace(str(path), path.name).replace(str(path.parent), "")
        raise UnreadableDocument(f"{path.name}: could not be read ({detail}).") from exc

    # Centrally, not per loader: a PDF and a Word file carry zero-width and
    # control characters just as a text file does, and cleaning in one branch
    # only is how three of the four formats keep the defect.
    text = _clean(text)

    if not text.strip():
        raise UnreadableDocument(f"{path.name}: empty.")
    return RawDocument(source=path.relative_to(root).as_posix(), text=text)


def load_corpus(
    root: str | Path, include: tuple[str, ...], encoding: str | None = None
) -> tuple[list[RawDocument], list[str]]:
    """Load every included document under `root`.

    Returns `(documents, skipped)` — skipped entries are human-readable reasons,
    surfaced by the CLI at the end of an ingest. **One bad file must not fail
    the run**: a corpus is a folder a client assembled by hand, and a single
    scanned fax in it is not a reason to deliver nothing.

    Documents come back sorted by source so that the corpus fingerprint does not
    depend on filesystem iteration order.
    """
    root_path = Path(root)
    if not root_path.is_dir():
        raise FileNotFoundError(f"Corpus directory not found: {root_path}")

    documents: list[RawDocument] = []
    skipped: list[str] = []
    suffixes = {ext.lower() for ext in include}
    # Resolved once: every file's real location is checked against this.
    root_real = root_path.resolve()

    for path in sorted(root_path.rglob("*")):
        if not path.is_file():
            continue

        if not _is_inside(path, root_real):
            # ⚠️ A symlink out of the corpus, and it is a privacy defect rather
            # than a tidiness one. Found by pointing `escaped.txt` at
            # `/etc/passwd` in a test corpus: it was read, chunked, embedded and
            # became searchable content — in the product whose entire promise is
            # that the client controls which documents are exposed.
            #
            # `is_file()` follows symlinks, so the check has to be explicit.
            # Reported rather than ignored: a link the client put there on
            # purpose (a network share, a folder of policies kept elsewhere) is
            # a reasonable thing to have done, and they need to be told it was
            # not included instead of wondering why those answers are missing.
            reason = f"{path.name}: points outside the documents folder; not included."
            logger.warning("skipping document: %s", reason)
            skipped.append(reason)
            continue

        if path.suffix.lower() not in suffixes:
            if path.name not in _CLUTTER and not path.name.startswith("~$"):
                # Silence here generated support email: a client who drops in a
                # spreadsheet and is told nothing concludes the assistant read
                # it, then reports "it doesn't know what's in my file" weeks
                # later. Naming it at ingest costs one line.
                reason = f"{path.name}: unsupported file type; not included."
                skipped.append(reason)
            continue

        try:
            documents.append(load_document(path, root_path, encoding))
        except UnreadableDocument as exc:
            logger.warning("skipping document: %s", exc)
            skipped.append(str(exc))

    return documents, skipped
